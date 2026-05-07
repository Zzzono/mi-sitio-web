from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

def add_title(text):
    p = doc.add_heading(text, level=0)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

def add_h1(text):
    doc.add_heading(text, level=1)

def add_h2(text):
    doc.add_heading(text, level=2)

def add_h3(text):
    doc.add_heading(text, level=3)

def add_p(text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p

def add_result(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 100, 0)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1+len(rows), cols=len(headers))
    table.style = 'Light Shading Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            table.rows[r_idx+1].cells[c_idx].text = str(val)
    doc.add_paragraph()

# ============ PORTADA ============
doc.add_paragraph()
doc.add_paragraph()
add_title("EJERCICIOS DE INGENIERÍA INDUSTRIAL")
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Diagramas de Actividades Múltiples, Hombre-Máquina y Punto de Equilibrio")
run.font.size = Pt(14)
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Soluciones Completas")
run.font.size = Pt(12)
run.italic = True
doc.add_page_break()

# ============ EJERCICIO 1 ============
add_h1("EJERCICIO 1: Diagrama de Actividades Múltiples — 2 Máquinas")
add_h2("Datos del problema")
add_p("Un operario tiene asignado 2 máquinas. Jornada: 8 h/día, 5 días/semana.")
add_table(
    ["Operación", "Tiempo (min)"],
    [
        ["Cargar la pieza en máquina", "0.80"],
        ["Inspeccionar", "0.10"],
        ["Maquinado automático", "2.60"],
        ["Traslado entre máquinas", "0.05"],
        ["Descargar la pieza", "0.50"],
    ]
)

add_h2("Desarrollo")
add_p("Tiempo del operario por máquina (To):")
add_p("To = Descargar + Cargar + Inspeccionar = 0.50 + 0.80 + 0.10 = 1.40 min", bold=True)
add_p("Tiempo automático de máquina (Tm): 2.60 min")
add_p("Traslado (Tt): 0.05 min")
add_p("")
add_p("Ciclo total del operario para 2 máquinas:")
add_p("= (To + Tt) × 2 = (1.40 + 0.05) × 2 = 2.90 min")
add_p("Ciclo de cada máquina:")
add_p("= To + Tm = 1.40 + 2.60 = 4.00 min")
add_result("Tiempo de ciclo = máx(2.90, 4.00) = 4.00 min (la máquina determina el ciclo)")
add_p("Tiempo de espera del operario = 4.00 − 2.90 = 1.10 min")

add_h2("Diagrama de Actividades Múltiples")
add_table(
    ["Tiempo (min)", "OPERARIO", "MÁQUINA 1", "MÁQUINA 2"],
    [
        ["0.00", "Descargar M1", "Descargando", "Maq. Automático"],
        ["0.50", "Cargar M1", "Cargando", "↓"],
        ["1.30", "Inspeccionar M1", "Inspección", "↓"],
        ["1.40", "Traslado → M2", "MAQ. AUTOMÁTICO", "↓"],
        ["1.45", "Descargar M2", "↓", "Descargando"],
        ["1.95", "Cargar M2", "↓", "Cargando"],
        ["2.75", "Inspeccionar M2", "↓", "Inspección"],
        ["2.85", "Traslado → M1", "↓", "MAQ. AUTOMÁTICO"],
        ["2.90", "★ ESPERA ★", "↓", "↓"],
        ["4.00", "Descargar M1 (nuevo ciclo)", "Descargando", "↓"],
    ]
)

add_h2("Cálculos de Producción")
add_table(
    ["Concepto", "Valor"],
    [
        ["Producción por ciclo", "2 piezas (1 por máquina)"],
        ["Tiempo de ciclo", "4.00 min"],
        ["Producción/hora", "2 × (60/4.00) = 30 piezas/hora"],
        ["Producción/día (8h)", "30 × 8 = 240 piezas/día"],
        ["Producción semanal (5d)", "240 × 5 = 1,200 piezas/semana"],
    ]
)

add_h2("Eficiencias")
add_table(
    ["Indicador", "Cálculo", "Resultado"],
    [
        ["% Eficiencia operario", "(2.90/4.00)×100", "72.50%"],
        ["% Utilización M1", "(2.60/4.00)×100", "65.00%"],
        ["% Utilización M2", "(2.60/4.00)×100", "65.00%"],
    ]
)
doc.add_page_break()

# ============ EJERCICIO 2 ============
add_h1("EJERCICIO 2: Diagrama de Actividades Múltiples — 3 Máquinas")
add_h2("Datos del problema")
add_p("Jornada: 8 h/día, 5 días/semana.")
add_table(
    ["Operación", "Tiempo (min)"],
    [
        ["Traslado entre máquinas", "0.80"],
        ["Cargar máquina", "0.50"],
        ["Descargar máquina", "1.00"],
        ["Tiempo automático", "5.00"],
        ["Embalar", "1.00"],
    ]
)

add_h2("Desarrollo")
add_p("Tiempo del operario por máquina (To):")
add_p("To = Descargar + Embalar + Cargar = 1.00 + 1.00 + 0.50 = 2.50 min", bold=True)
add_p("Tiempo automático (Tm): 5.00 min")
add_p("Traslado (Tt): 0.80 min")
add_p("")
add_p("Ciclo total del operario para 3 máquinas:")
add_p("= (To + Tt) × 3 = (2.50 + 0.80) × 3 = 9.90 min")
add_p("Ciclo de cada máquina:")
add_p("= To + Tm = 2.50 + 5.00 = 7.50 min")
add_result("Tiempo de ciclo = máx(9.90, 7.50) = 9.90 min (el operario es el cuello de botella)")
add_p("Tiempo de espera de cada máquina = 9.90 − 7.50 = 2.40 min")

add_h2("Diagrama de Actividades Múltiples")
add_table(
    ["Tiempo (min)", "OPERARIO", "MÁQUINA 1", "MÁQUINA 2", "MÁQUINA 3"],
    [
        ["0.00", "Descargar M1", "Descargando", "Maq. Autom.", "Maq. Autom."],
        ["1.00", "Embalar M1", "Embalando", "↓", "↓"],
        ["2.00", "Cargar M1", "Cargando", "↓", "↓"],
        ["2.50", "Traslado → M2", "MAQ. AUTOM.", "↓", "↓"],
        ["3.30", "Descargar M2", "↓", "Descargando", "↓"],
        ["4.30", "Embalar M2", "↓", "Embalando", "↓"],
        ["4.80", "Cargar M2", "↓", "Cargando", "↓"],
        ["5.80", "Traslado → M3", "↓", "MAQ. AUTOM.", "↓"],
        ["6.60", "Descargar M3", "↓", "↓", "Descargando"],
        ["7.50", "(M1 termina)", "★ ESPERA ★", "↓", "↓"],
        ["7.60", "Embalar M3", "↓", "↓", "Embalando"],
        ["8.60", "Cargar M3", "↓", "↓", "Cargando"],
        ["9.10", "Traslado → M1", "↓", "↓", "MAQ. AUTOM."],
        ["9.90", "Descargar M1", "Descargando", "↓", "↓"],
    ]
)

add_h2("Cálculos de Producción")
add_table(
    ["Concepto", "Valor"],
    [
        ["Producción por ciclo", "3 piezas (1 por máquina)"],
        ["Tiempo de ciclo", "9.90 min"],
        ["Producción/hora", "3 × (60/9.90) = 18.18 piezas/hora"],
        ["Producción/día (8h)", "18.18 × 8 = 145.45 piezas/día"],
        ["Producción semanal (5d)", "145.45 × 5 ≈ 727 piezas/semana"],
    ]
)

add_h2("Eficiencias")
add_table(
    ["Indicador", "Cálculo", "Resultado"],
    [
        ["% Eficiencia operario", "(9.90/9.90)×100", "100% (sin ocio)"],
        ["% Utilización cada máquina", "(5.00/9.90)×100", "50.51%"],
    ]
)
doc.add_page_break()

# ============ EJERCICIO 3 ============
add_h1("EJERCICIO 3: Diagrama HOMBRE–MÁQUINA")
add_h2("Datos del problema")
add_p("Capacidad de la máquina: 150 KG")
add_table(
    ["Operación", "Tiempo (min)"],
    [
        ["Cargar material", "15"],
        ["Encender máquina", "5"],
        ["Inspección inicial (durante maquinado)", "5"],
        ["Maquinado (automático)", "110"],
        ["Descargar", "10"],
    ]
)
add_table(
    ["Costo", "Valor"],
    [
        ["Costo hora-máquina", "S/. 5.00"],
        ["Costo hora-hombre", "S/. 30.00"],
        ["Costo material por KG procesado", "S/. 3.50"],
    ]
)

add_h2("Diagrama HOMBRE–MÁQUINA")
add_table(
    ["Tiempo (min)", "OPERARIO", "MÁQUINA"],
    [
        ["0 – 15", "Cargar material (TRABAJA)", "Cargando (INACTIVA)"],
        ["15 – 20", "Encender máquina (TRABAJA)", "Encendiendo (INACTIVA)"],
        ["20 – 25", "Inspección inicial (TRABAJA)", "Maquinado automático (TRABAJA)"],
        ["25 – 130", "OCIOSO (105 min)", "Maquinado automático (TRABAJA)"],
        ["130 – 140", "Descargar (TRABAJA)", "Descargando (INACTIVA)"],
    ]
)

add_h2("Soluciones")

add_h3("a) Tiempo de ciclo")
add_p("TC = Cargar + Encender + Maquinado + Descargar")
add_p("TC = 15 + 5 + 110 + 10")
add_result("Tiempo de ciclo = 140 minutos")

add_h3("b) Tiempo de trabajo del operario")
add_p("T_operario = Cargar + Encender + Inspección + Descargar")
add_p("T_operario = 15 + 5 + 5 + 10")
add_result("Tiempo de trabajo del operario = 35 minutos")

add_h3("c) Tiempo de trabajo de la máquina")
add_p("T_máquina = Maquinado automático")
add_result("Tiempo de trabajo de la máquina = 110 minutos")

add_h3("d) % de Eficiencia del operario")
add_p("% Efic. = (T_trabajo_operario / T_ciclo) × 100")
add_p("% Efic. = (35 / 140) × 100")
add_result("Eficiencia del operario = 25.00%")

add_h3("e) % de Utilización de la máquina")
add_p("% Util. = (T_trabajo_máquina / T_ciclo) × 100")
add_p("% Util. = (110 / 140) × 100")
add_result("Utilización de la máquina = 78.57%")

add_h3("f) Producción por hora")
add_p("Producción/ciclo = 150 kg")
add_p("Ciclos/hora = 60 / 140 = 0.4286 ciclos/hora")
add_p("Producción/hora = 0.4286 × 150")
add_result("Producción por hora = 64.29 kg/hora")

add_h3("g) Capacidad de atención del operario")
add_p("N = (Tm + To) / To")
add_p("N = (110 + 35) / 35 = 140 / 35")
add_result("El operario puede atender hasta 4 máquinas")

add_h3("h) Costos totales de operación (por ciclo)")
add_table(
    ["Concepto", "Cálculo", "Valor"],
    [
        ["Costo hora-máquina", "S/.5.00/hr × (140/60)hr", "S/. 11.67"],
        ["Costo hora-hombre", "S/.30.00/hr × (140/60)hr", "S/. 70.00"],
        ["Costo material", "S/.3.50/kg × 150 kg", "S/. 525.00"],
        ["TOTAL", "", "S/. 606.67"],
    ]
)
add_result("Costo total de operación por ciclo = S/. 606.67")

add_h3("i) Productividad total (kg/S/.)")
add_p("Productividad = Producción / Costo total")
add_p("Productividad = 150 / 606.67")
add_result("Productividad total = 0.2473 kg/S/.")

add_h2("Resumen Ejercicio 3")
add_table(
    ["Inciso", "Concepto", "Resultado"],
    [
        ["a)", "Tiempo de ciclo", "140 min"],
        ["b)", "Tiempo trabajo operario", "35 min"],
        ["c)", "Tiempo trabajo máquina", "110 min"],
        ["d)", "% Eficiencia operario", "25.00%"],
        ["e)", "% Utilización máquina", "78.57%"],
        ["f)", "Producción por hora", "64.29 kg/hr"],
        ["g)", "Capacidad de atención", "4 máquinas"],
        ["h)", "Costo total operación", "S/. 606.67"],
        ["i)", "Productividad total", "0.2473 kg/S/."],
    ]
)
doc.add_page_break()

# ============ EJERCICIO 4 ============
add_h1("EJERCICIO 4: Diagrama de Actividades Múltiples — 2 Máquinas")
add_h2("Datos del problema")
add_table(
    ["Operación", "Tiempo (min)"],
    [
        ["Traslado entre máquinas", "0.25"],
        ["Cargar máquina", "1.00"],
        ["Descargar máquina", "0.50"],
        ["Tiempo automático", "5.50"],
        ["Alistar material", "2.00"],
        ["Inspeccionar", "0.75"],
    ]
)

add_h2("Desarrollo")
add_p("Tiempo del operario por máquina (To):")
add_p("To = Descargar + Inspeccionar + Alistar + Cargar = 0.50 + 0.75 + 2.00 + 1.00 = 4.25 min", bold=True)
add_p("Tiempo automático (Tm): 5.50 min")
add_p("Traslado (Tt): 0.25 min")
add_p("")
add_p("Ciclo total del operario para 2 máquinas:")
add_p("= (To + Tt) × 2 = (4.25 + 0.25) × 2 = 9.00 min")
add_p("Ciclo de cada máquina:")
add_p("= To + Tm = 4.25 + 5.50 = 9.75 min")
add_result("Tiempo de ciclo = máx(9.00, 9.75) = 9.75 min (la máquina determina el ciclo)")
add_p("Tiempo de espera del operario = 9.75 − 9.00 = 0.75 min")

add_h2("Diagrama de Actividades Múltiples")
add_table(
    ["Tiempo (min)", "OPERARIO", "MÁQUINA 1", "MÁQUINA 2"],
    [
        ["0.00", "Descargar M1", "Descargando", "Maq. Automático"],
        ["0.50", "Inspeccionar M1", "Inspeccionando", "↓"],
        ["1.25", "Alistar mat. M1", "Alistando", "↓"],
        ["3.25", "Cargar M1", "Cargando", "↓"],
        ["4.25", "Traslado → M2", "MAQ. AUTOMÁTICO", "↓"],
        ["4.50", "Descargar M2", "↓", "Descargando"],
        ["5.00", "Inspeccionar M2", "↓", "Inspeccionando"],
        ["5.75", "Alistar mat. M2", "↓", "Alistando"],
        ["7.75", "Cargar M2", "↓", "Cargando"],
        ["8.75", "Traslado → M1", "↓", "MAQ. AUTOMÁTICO"],
        ["9.00", "★ ESPERA ★", "↓", "↓"],
        ["9.75", "Descargar M1 (nuevo ciclo)", "Descargando", "↓"],
    ]
)

add_h2("Cálculos")
add_table(
    ["Concepto", "Valor"],
    [
        ["Producción por ciclo", "2 piezas"],
        ["Tiempo de ciclo", "9.75 min"],
        ["Producción/hora", "2 × (60/9.75) = 12.31 piezas/hora"],
    ]
)
add_table(
    ["Indicador", "Cálculo", "Resultado"],
    [
        ["% Eficiencia operario", "(9.00/9.75)×100", "92.31%"],
        ["% Utilización cada máquina", "(5.50/9.75)×100", "56.41%"],
    ]
)
doc.add_page_break()

# ============ EJERCICIO 5 ============
add_h1("EJERCICIO 5: Punto de Equilibrio — Alquiler de Montacargas")
add_h2("Datos del problema")
add_table(
    ["Concepto", "Valor"],
    [
        ["Ventas mensuales", "S/. 480,000"],
        ["Servicios mensuales", "4,800 servicios"],
        ["Costos Variables Totales (CVT)", "S/. 360,000"],
        ["Costos Fijos (CF)", "S/. 90,000"],
        ["Montacargas disponibles", "24"],
    ]
)

add_h2("Desarrollo")

add_h3("Paso 1: Precio de venta unitario (PVu)")
add_p("PVu = Ventas / N° servicios = 480,000 / 4,800")
add_result("PVu = S/. 100.00 por servicio")

add_h3("Paso 2: Costo variable unitario (CVu)")
add_p("CVu = CVT / N° servicios = 360,000 / 4,800")
add_result("CVu = S/. 75.00 por servicio")

add_h3("Paso 3: Margen de contribución unitario (MCu)")
add_p("MCu = PVu − CVu = 100.00 − 75.00")
add_result("MCu = S/. 25.00 por servicio")

add_h3("Paso 4: Punto de equilibrio en servicios")
add_p("P.E. = CF / MCu = 90,000 / 25.00")
add_result("P.E. = 3,600 servicios/mes")

add_h3("Paso 5: Punto de equilibrio en soles")
add_p("P.E. (S/.) = 3,600 × 100")
add_result("P.E. = S/. 360,000")

add_h3("Paso 6: Servicios por montacargas")
add_p("Servicios/montacargas = 4,800 / 24 = 200 servicios/montacargas/mes")

add_h3("Paso 7: Montacargas mínimos para el P.E.")
add_p("Montacargas mínimos = P.E. (servicios) / Servicios por montacargas")
add_p("Montacargas mínimos = 3,600 / 200")
add_result("LA EMPRESA DEBE ALQUILAR UN MÍNIMO DE 18 MONTACARGAS PARA ALCANZAR EL PUNTO DE EQUILIBRIO")

add_h2("Verificación")
add_table(
    ["Concepto", "Valor"],
    [
        ["Ingresos P.E.", "18 × 200 × S/.100 = S/. 360,000"],
        ["CVT en P.E.", "18 × 200 × S/.75 = S/. 270,000"],
        ["Costos Fijos", "S/. 90,000"],
        ["Costo Total", "S/. 360,000"],
        ["Utilidad", "360,000 − 360,000 = S/. 0 ✓"],
    ]
)

# ============ GUARDAR ============
output_path = os.path.join(os.path.expanduser("~"), "Downloads", "Ejercicios_Ingenieria_Industrial.docx")
doc.save(output_path)
print(f"Documento guardado en: {output_path}")
